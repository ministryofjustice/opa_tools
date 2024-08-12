import argparse
import os
import sys

import docx


class Docx2Ruletxt:
    def __init__(self, input_filepath, output_filepath=None):
        '''Converts a docx file into a ruletxt file'''
        if os.path.basename(input_filepath).startswith('~$'):
            if output_filepath:
                print(f'Ignoring tempfile created by Word {input_filepath}')
            return

        self.output_file = open(output_filepath, 'w') if output_filepath else None

        document = docx.Document(input_filepath)
        self.current_indent = 1
        for doc_index, doc_object in enumerate(document.iter_inner_content()):
            if isinstance(doc_object, docx.text.paragraph.Paragraph):
                self.process_paragraph(doc_object, doc_index)
            elif isinstance(doc_object, docx.table.Table):
                if self.current_indent > 1:
                    self.close_paragraph_brackets()
                self.process_table(doc_object, doc_index)
            else:
                self.print_('ERROR', f'Unhandled doc object: {type(doc_object)}')
        
        if output_filepath:
            print(f"Processed {input_filepath} -> {output_filepath}")

    def print_(self, header_tag, body):
        padding = ' ' * (max(len('[table-OPM-conclusion]') - len(header_tag) - 2, 0))
        body = body.rstrip()
        if '\n' in body:
            # very rare to see a newline inside a paragraph or table cell.
            # print a symbol, rather than have lines without a header tag
            body = body.replace('\n', '\\n')
        line = f'[{header_tag}] {padding} {body}'
        if self.output_file:
            print(line.rstrip(), file=self.output_file)
        else:
            print(line)

    def close_paragraph_brackets(self):
        self.print_('close-bracket' + ')' * (self.current_indent - 1), '')
        self.current_indent = 1

    def process_paragraph(self, p, doc_index):
        style = p.style.style_id
        indent = 1

        if style.startswith('OPM-level'):
            # add brackets, to make it parsable with a CFL
            indent = int(style.split('OPM-level')[1])
            if indent > self.current_indent:
                style += '(' * (indent - self.current_indent)
            elif indent < self.current_indent:
                style += ')' * (self.current_indent - indent)
            self.current_indent = indent
        if self.current_indent > 1 and style == 'OPM-conclusion':
            self.close_paragraph_brackets()

        self.print_(style, '    ' * (indent - 1)+ f'{p.text}')

    def process_table(self, table, doc_index):
        for row_index, r in enumerate(table.rows):
            row_cells = tuple(c for c in r.cells)

            if row_index == 0 and row_cells[0].paragraphs[0].style.style_id != 'OPM-conclusion':
                # this table isn't a rule
                return self.process_nonrule_table(table, doc_index)

            # hack to fix one table in 'Work Package 3/WP3 Calculation 30 -37.docx'
            if len(row_cells) == 3 and row_cells[0].text == row_cells[1].text:
                row_cells = (row_cells[0], row_cells[2])

            if len(row_cells) != 2:
                self.print_('ERROR', f'table row has too many cells {doc_index=} {row_index=} {len(row_cells)=}')
                breakpoint()
                continue

            cell1_style = row_cells[1].paragraphs[0].style.style_id
            # We're only interested in the styles in the 2nd column - the 1st column is always OPM-conclusion
            if row_index == 0:
                # In Word, this first row is a single merged cell, but it comes through to here as two duplicate cells
                self.print_(f'table-{cell1_style}', f'{row_cells[1].paragraphs[0].text}')
            else:
                for i, paragraph in enumerate(row_cells[1].paragraphs):
                    if i == 0:
                        self.print_(f'table-{cell1_style}', f'{row_cells[0].text} | [1] {paragraph.text}')
                    else:
                        indent = 0
                        if paragraph.style.style_id.startswith('OPM-level'):
                            indent = int(paragraph.style.style_id.split('OPM-level')[1])
                        self.print_('table-cell-continued', f'  | [{indent}] ' + '    ' * (indent - 1) + f'{paragraph.text}')

        # self.print_('table-end', '')

    def process_nonrule_table(self, table, doc_index):
        for row_index, r in enumerate(table.rows):
            row_cells = tuple(c for c in r.cells)
            self.print_('nonrule-table', ' | '.join((cell.text for cell in row_cells)))

def convert_filename(filename):
    return filename.replace(' ', '_').replace('.docx', '.ruletxt')

def convert_directory_path(relative_path):
    return relative_path.replace(' ', '_')

def process_directory(input_dir, output_dir):
    # Process all files in the input directory recursively
    for root, _, files in os.walk(input_dir):
        for file in files:
            if not file.endswith('.docx'):
                continue
            input_file = os.path.join(root, file)
            relative_path = convert_directory_path(os.path.relpath(root, input_dir))
            output_file = os.path.join(output_dir, relative_path, convert_filename(file))

            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            Docx2Ruletxt(input_file, output_file)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert files or directories from docx to ruletxt format.")

    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('-i', '--input', help="Input file (docx format)")
    group.add_argument('-d', '--input-dir', help="Input directory (containing docs files)")

    parser.add_argument('-o', '--output', help="Output file or directory")

    args = parser.parse_args()

    if args.input:
        Docx2Ruletxt(args.input, args.output)
    elif args.input_dir:
        process_directory(args.input_dir, args.output)
