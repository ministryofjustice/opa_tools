import os
import sys

from docx import Document


#input_filepath = os.path.expanduser("~/code/laa-ccms-opa-policy-models-zips-extracted/MeansAssessment/Rules/LAR/LAR System Rules.docx")
if len(sys.argv) > 1:
    input_filepath = sys.argv[1]
else:
    print("No input filepath provided")
    sys.exit(1)

if len(sys.argv) > 2:
    out_filepath = sys.argv[2]
    f = open(out_filepath, 'w')
    sys.stdout = f

document = Document(input_filepath)

# Paragraphs

# Properties:
# text: someone receives a passported benefit and
# alignment: None
# contains_page_break: False
# hyperlinks: []
# paragraph_format: <docx.text.parfmt.ParagraphFormat object at 0x107d03f20>
# part: <docx.parts.document.DocumentPart object at 0x107c8d8e0>
# rendered_page_breaks: []
# runs: [<docx.text.run.Run object at 0x107d03cb0>, <docx.text.run.Run object at 0x107d03bc0>]
# style: _ParagraphStyle('OPM - level 2') id: 4426054592

# Methods:
# clear
# insert_paragraph_before
# iter_inner_content

# Style
# names and IDs:
# 'OPM - conclusion'  'OPM-conclusion'
# 'OPM - level 1'     'OPM-level1'
# etc

parse_state = None
logic = []
code = []

def finish_the_if_block(code, conclusion_attribute, logic):
    code.append(f'{conclusion_attribute} = \\')
    for level, logic_line in logic:
        code.append('    ' * level + f'{logic_line}')
    code.append('')
    print('\n'.join(code))
    # clean up
    code.clear()
    logic.clear()

for paragraph_index, p in enumerate(document.paragraphs):
    style = p.style.style_id
    if style == 'OPM-conclusion':
        if parse_state:
            # the end of previous 'if block'
            finish_the_if_block(code, conclusion_attribute, logic)
            parse_state = None

        parse_state = 'if block'
        conclusion_attribute = p.text.split(' if')[0]
    elif style.startswith('OPM-level'):
        if parse_state != 'if block':
            print(f'ERROR, logic outside a block paragraph={paragraph_index} {style=} text={p.text}')
            break
        level = int(style.split('OPM-level')[1])
        logic.append((level, p.text))
    elif style == 'OPM-blankline':
        print('\n')
    elif style in ('OPM-commentary', 'OPM-RuleName') or style.startswith('OPM-Heading') or style.startswith('TOC') or style.startswith('Comment-'):
        if p.text.strip():
            lines = p.text.split('\n')
            for line in lines:
                print(f'# {line}')
            print('\n')
    else:
        print(f'IGNORE paragraph paragraph={paragraph_index} {style=} text={p.text}')

if parse_state == 'if block':
    finish_the_if_block(code, conclusion_attribute, logic)
    
    # "if" conclusion
    # 80 the passported test is conducted on the carer or the carer's partner if - style: OPM-conclusion
    # 81 both - style: OPM-level1
    # 82 the carer is included in the assessment and - style: OPM-level2
    # 83 the carer receives the passported benefit - style: OPM-level2
    # 88  - style: OPM-blankline

    # "Equals" conclusion
    # 33 the will name change LSC to LAA = the name change LSC to LAA - style: OPM-conclusion
    # 34  - style: OPM-blankline


# Tables

# Properties
# autofit: True
# columns: <docx.table._Columns object at 0x106c97a70>
# part: <docx.parts.document.DocumentPart object at 0x106c1c200>
# rows: <docx.table._Rows object at 0x106c97b30>
# style: _TableStyle('Normal Table') id: 4408835392
# table_direction: None

# Methods:
# add_column
# add_row
# cell
# column_cells
# row_cells

# Styles example:
# the name change Funding Code to LAPSO
# cell0_style='OPM-conclusion' cell1_style='OPM-conclusion'
# "Funding Code" | the LAR rules do not apply to this application
# cell0_style='OPM-conclusion' cell1_style='OPM-level1'
# "Lord Chancellor’s Guidance on financial eligibility for certificated work" |	otherwise
# cell0_style='OPM-conclusion' cell1_style='OPM-Alternativeconclusion'

def parse_table(table, table_index):
    code = []
    for row_index, r in enumerate(table.rows):
        row_cells = tuple(c for c in r.cells)
        if len(row_cells) != 2:
            print(f'ERROR, table row has too many cells {table_index=} {row_index=} {len(row_cells)=}')
            return None
        
        cell0_style = row_cells[0].paragraphs[0].style.style_id
        cell1_style = row_cells[1].paragraphs[0].style.style_id    
        # print (f'{cell0_style=} {cell1_style=}')
        if row_index == 0:
            if cell1_style != 'OPM-conclusion':
                print(f'ERROR, table first row should be style: OPM-conclusion {table_index=} {row_index=} {cell0_style=}')
                return None
            conclusion_attribute = row_cells[1].paragraphs[0].text
            code.append(f'{conclusion_attribute} = ')
        elif cell1_style == 'OPM-level1':
            value = row_cells[0].paragraphs[0].text
            condition = row_cells[1].paragraphs[0].text
            if row_index != 1:
                code[-1] += f' else\\'
            code.append(f'    {value} if ({condition})')
        elif cell1_style == 'OPM-Alternativeconclusion':
            value = row_cells[0].paragraphs[0].text
            if row_index != 1:
                code[-1] += f' else\\'
            code.append(f'    {value}')            
        else:
            print(f'ERROR, table unknown style: {table_index=} {row_index=} {cell1_style=}')
            return None
    return code
                
for table_index, table in enumerate(document.tables):
    # print(f'Table {table_index}:')
    code = parse_table(table, table_index)
    print('\n'.join(code))
    print('\n')
    